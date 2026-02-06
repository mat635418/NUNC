import streamlit as st
import openai
from docx import Document
from redlines import Redlines
import io

# --- CONFIGURAZIONE PAGINA & BRANDING ---
st.set_page_config(
    page_title="NUNC - Legal Updater", 
    layout="wide", 
    page_icon="⚖️"
)

# --- STILE NUNC & LOGO BB (CSS) ---
st.markdown("""
<style>
    /* Colori Istituzionali: Navy Blue e Oro Antico */
    :root {
        --primary-color: #0f172a;
        --accent-color: #d4af37;
        --text-color: #334155;
    }
    
    /* Header Container */
    .nunc-header {
        display: flex;
        align-items: center;
        justify-content: center;
        padding: 2rem;
        background-color: var(--primary-color);
        border-radius: 12px;
        margin-bottom: 2rem;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
    }
    
    /* Logo BB Costruito in CSS */
    .bb-logo {
        font-family: 'Times New Roman', serif;
        font-weight: 700;
        font-size: 3.5rem;
        color: var(--accent-color);
        border: 3px solid var(--accent-color);
        padding: 5px 15px;
        border-radius: 8px;
        margin-right: 20px;
        letter-spacing: -5px; /* Sovrappone leggermente le lettere */
        background: rgba(255,255,255,0.05);
    }
    
    /* Nome App */
    .app-name {
        font-family: 'Helvetica Neue', sans-serif;
        font-weight: 300;
        font-size: 3rem;
        color: white;
        letter-spacing: 8px;
        text-transform: uppercase;
    }
    
    /* Payoff */
    .payoff {
        color: #94a3b8;
        font-size: 0.9rem;
        text-transform: uppercase;
        letter-spacing: 2px;
        margin-top: -10px;
        text-align: center;
    }

    /* Stile Bottoni */
    div.stButton > button {
        background-color: var(--primary-color);
        color: white;
        border: 1px solid var(--accent-color);
        font-weight: bold;
    }
    div.stButton > button:hover {
        background-color: var(--accent-color);
        color: var(--primary-color);
        border-color: var(--primary-color);
    }
</style>
""", unsafe_allow_html=True)

# --- MOTORE DI ELABORAZIONE ---

def read_docx(file):
    """Legge il contenuto del file Word caricato."""
    doc = Document(file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def create_docx(text):
    """Crea un file Word scaricabile dal testo elaborato."""
    doc = Document()
    doc.add_heading('Aggiornamento NUNC', 0)
    for paragraph in text.split('\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def nunc_engine(testo_vigente, testo_novita, api_key):
    """
    Motore NUNC: Aggiornamento in tempo reale.
    """
    client = openai.OpenAI(api_key=api_key)
    
    prompt_sistema = """
    Sei NUNC, un sistema esperto di aggiornamento normativo.
    Il tuo compito è armonizzare un TESTO VIGENTE con una NOVITÀ NORMATIVA.
    
    DIRETTIVE:
    1. Precisione Assoluta: Modifica date, importi e requisiti esattamente come indicato nella novità.
    2. Conservazione: Non toccare le parti del testo vigente che non sono impattate dalla novità.
    3. Stile: Mantieni il tono professionale, asettico e giuridico del testo originale.
    4. Output: Restituisci solo il testo finale completo.
    """
    
    user_msg = f"""
    --- TESTO DA AGGIORNARE (DATABASE) ---
    {testo_vigente}
    
    --- NUOVA DISPOSIZIONE (AGGIORNAMENTO) ---
    {testo_novita}
    """

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": prompt_sistema},
                {"role": "user", "content": user_msg}
            ],
            temperature=0.1
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Errore NUNC: {str(e)}"

# --- HEADER GRAFICO ---
st.markdown("""
<div class="nunc-header">
    <div class="bb-logo">BB</div>
    <div>
        <div class="app-name">NUNC</div>
    </div>
</div>
<div class="payoff">Normativa Aggiornata in Tempo Reale</div>
<br>
""", unsafe_allow_html=True)

# --- SIDEBAR CONFIG ---
with st.sidebar:
    st.header("🔐 Accesso Sistema")
    api_key = st.text_input("API Key", type="password")
    st.markdown("---")
    st.info("**Modalità Operativa:**\nRecepimento automatico variazioni legislative su testi proprietari.")

# --- AREA DI LAVORO ---
col_in1, col_in2 = st.columns(2)

with col_in1:
    st.subheader("📂 Banca Dati (Word)")
    uploaded_file = st.file_uploader("Carica documento vigente", type=['docx'])
    testo_vigente = ""
    if uploaded_file:
        testo_vigente = read_docx(uploaded_file)
        st.caption(f"Letto documento di {len(testo_vigente)} caratteri.")

with col_in2:
    st.subheader("⚖️ Novità (Testo)")
    testo_novita = st.text_area(
        "Incolla circolare o legge", 
        height=200,
        placeholder="Es: A decorrere dal 1° Gennaio, l'aliquota passa al..."
    )

# --- AZIONE ---
st.markdown("---")
col_center = st.columns([1, 2, 1])
with col_center[1]:
    start_btn = st.button("ESEGUI AGGIORNAMENTO NORMATIVO", use_container_width=True)

if start_btn:
    if not api_key or not uploaded_file or not testo_novita:
        st.error("Errore: Verificare API Key e documenti in ingresso.")
    else:
        with st.spinner("Analisi differenziale in corso..."):
            # Elaborazione
            testo_aggiornato = nunc_engine(testo_vigente, testo_novita, api_key)
            
            # Generazione Delta
            differenze = Redlines(testo_vigente, testo_aggiornato)
            
            st.success("Elaborazione completata.")
            
            # Output Visivo
            tab1, tab2 = st.tabs(["📝 Revisione (Track Changes)", "📄 Documento Finale"])
            
            with tab1:
                st.markdown("### Dettaglio Variazioni")
                st.markdown(differenze.output_markdown, unsafe_allow_html=True)
                
            with tab2:
                st.markdown("### Testo Pronto per Archivio")
                st.text_area("Anteprima", value=testo_aggiornato, height=300)
                
                doc_finale = create_docx(testo_aggiornato)
                st.download_button(
                    label="Scarica DOCX Aggiornato",
                    data=doc_finale,
                    file_name="NUNC_Update.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
